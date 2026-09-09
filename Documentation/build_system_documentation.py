from __future__ import annotations

import json
import os
import re
import textwrap
from datetime import datetime
from pathlib import Path

from PIL import Image, ImageDraw, ImageFont
from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DOC = ROOT / "Documentation"
OUT = Path(os.environ.get("SMARTFOUNDATION_DOCX_OUT", DOC / "SmartFoundation_System_Documentation.docx"))
PROGRESS = DOC / "Documentation-Progress.json"
FIG_DIR = DOC / ".word_qa" / "figures"

FONT = "Arial"
MONO = "Consolas"
NAVY, TEAL, GRAY = "183B56", "187B8C", "58656D"
LIGHT, PALE, GOLD, RED = "EAF3F5", "F5F7F8", "9A6B00", "A33A3A"
PAGE_WIDTH_DXA = 11906
CONTENT_DXA = 9866


def rtl(p, align=WD_ALIGN_PARAGRAPH.RIGHT):
    p.alignment = align
    ppr = p._p.get_or_add_pPr()
    node = ppr.find(qn("w:bidi"))
    if node is None:
        node = OxmlElement("w:bidi")
        ppr.append(node)
    node.set(qn("w:val"), "1")
    return p


def font(run, size=10.5, bold=None, color=None, name=FONT, italic=None):
    run.font.name = name
    rpr = run._element.get_or_add_rPr()
    rf = rpr.rFonts
    for attr in ("ascii", "hAnsi", "cs", "eastAsia"):
        rf.set(qn(f"w:{attr}"), name)
    if rpr.find(qn("w:rtl")) is None:
        rpr.append(OxmlElement("w:rtl"))
    run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_field(run, instruction, display=""):
    for kind, text in (("begin", None), (None, instruction), ("separate", None), (None, display), ("end", None)):
        if kind:
            el = OxmlElement("w:fldChar")
            el.set(qn("w:fldCharType"), kind)
        else:
            el = OxmlElement("w:instrText" if text == instruction else "w:t")
            el.set(qn("xml:space"), "preserve")
            el.text = text
        run._r.append(el)


def setup(doc):
    sec = doc.sections[0]
    sec.page_width, sec.page_height = Cm(21), Cm(29.7)
    sec.top_margin, sec.bottom_margin = Cm(1.75), Cm(1.65)
    sec.left_margin, sec.right_margin = Cm(1.8), Cm(1.8)
    sec.header_distance, sec.footer_distance = Cm(.65), Cm(.65)
    sec.different_first_page_header_footer = True
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name, normal.font.size = FONT, Pt(10.2)
    normal._element.rPr.rFonts.set(qn("w:cs"), FONT)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.16
    for name, size, color, before, after in (
        ("Title", 28, NAVY, 0, 10), ("Subtitle", 15, GRAY, 0, 9),
        ("Heading 1", 18, NAVY, 14, 7), ("Heading 2", 14, TEAL, 11, 5),
        ("Heading 3", 12, NAVY, 8, 4)):
        s = styles[name]
        s.font.name, s.font.size = FONT, Pt(size)
        s._element.rPr.rFonts.set(qn("w:cs"), FONT)
        s.font.color.rgb = RGBColor.from_string(color)
        s.font.bold = name != "Subtitle"
        s.paragraph_format.space_before, s.paragraph_format.space_after = Pt(before), Pt(after)
        s.paragraph_format.keep_with_next = True
    for name, size, color in (("Caption", 9, GRAY), ("Code RTL", 7.2, GRAY)):
        if name not in styles:
            s = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        else:
            s = styles[name]
        s.font.name, s.font.size = FONT if name == "Caption" else MONO, Pt(size)
        s.font.color.rgb = RGBColor.from_string(color)
        s.paragraph_format.space_after = Pt(5)
        s.paragraph_format.keep_with_next = name == "Caption"
    header = sec.header.paragraphs[0]
    rtl(header)
    font(header.add_run("SmartFoundation | التوثيق الفني والوظيفي"), 8.5, True, NAVY)
    footer = sec.footer.paragraphs[0]
    rtl(footer, WD_ALIGN_PARAGRAPH.CENTER)
    r = font(footer.add_run("الإصدار 1.0.0  |  2026-09-05  |  صفحة "), 8, color=GRAY)
    add_field(r, " PAGE ", "1")
    font(footer.add_run(" من "), 8, color=GRAY)
    add_field(footer.runs[-1], " NUMPAGES ", "1")


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    rtl(p)
    font(p.add_run(text), {1:18, 2:14, 3:12}[level], True, {1:NAVY, 2:TEAL, 3:NAVY}[level])
    return p


def para(doc, text="", size=10.2, bold=False, color=None, center=False):
    p = doc.add_paragraph()
    rtl(p, WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.RIGHT)
    font(p.add_run(text), size, bold, color)
    return p


def page_break(doc):
    doc.add_page_break()


def shade(cell, fill):
    tcpr = cell._tc.get_or_add_tcPr()
    shd = tcpr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tcpr.append(shd)
    shd.set(qn("w:fill"), fill)


def cell_margins(cell, top=85, start=120, bottom=85, end=120):
    tcpr = cell._tc.get_or_add_tcPr()
    mar = tcpr.find(qn("w:tcMar"))
    if mar is None:
        mar = OxmlElement("w:tcMar")
        tcpr.append(mar)
    for tag, val in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        n = mar.find(qn(f"w:{tag}"))
        if n is None:
            n = OxmlElement(f"w:{tag}")
            mar.append(n)
        n.set(qn("w:w"), str(val)); n.set(qn("w:type"), "dxa")


def geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    pr = table._tbl.tblPr
    bidi = pr.find(qn("w:bidiVisual"))
    if bidi is None:
        bidi = OxmlElement("w:bidiVisual")
        pr.append(bidi)
    bidi.set(qn("w:val"), "1")
    layout = pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tw = pr.find(qn("w:tblW")); tw.set(qn("w:w"), str(sum(widths))); tw.set(qn("w:type"), "dxa")
    ind = pr.find(qn("w:tblInd"))
    if ind is None:
        ind = OxmlElement("w:tblInd")
        pr.append(ind)
    ind.set(qn("w:w"), "120"); ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for c in list(grid): grid.remove(c)
    for width in widths:
        c = OxmlElement("w:gridCol"); c.set(qn("w:w"), str(width)); grid.append(c)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tcpr = cell._tc.get_or_add_tcPr(); tcw = tcpr.find(qn("w:tcW"))
            if tcw is None:
                tcw = OxmlElement("w:tcW")
                tcpr.append(tcw)
            tcw.set(qn("w:w"), str(widths[i])); tcw.set(qn("w:type"), "dxa")
            cell_margins(cell); cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    trpr = table.rows[0]._tr.get_or_add_trPr(); rep = OxmlElement("w:tblHeader"); rep.set(qn("w:val"), "true"); trpr.append(rep)


def prevent_row_split(row):
    """Keep a table row intact across page boundaries."""
    trpr = row._tr.get_or_add_trPr()
    node = trpr.find(qn("w:cantSplit"))
    if node is None:
        node = OxmlElement("w:cantSplit")
        trpr.append(node)
    node.set(qn("w:val"), "true")


def keep_table_block(table, keep_whole=False):
    """Prevent orphan headers and, for compact diagrams, keep the block together."""
    for row_index, row in enumerate(table.rows):
        prevent_row_split(row)
        keep_next = row_index == 0 or (keep_whole and row_index < len(table.rows) - 1)
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                paragraph.paragraph_format.keep_together = True
                paragraph.paragraph_format.keep_with_next = keep_next


def table(doc, headers, rows, widths=None, size=8.0, caption=None):
    if caption:
        p = doc.add_paragraph(style="Caption"); rtl(p); font(p.add_run(caption), 9, True, GRAY)
    t = doc.add_table(rows=1, cols=len(headers)); t.style = "Table Grid"
    if widths is None:
        widths = [CONTENT_DXA // len(headers)] * len(headers); widths[-1] += CONTENT_DXA - sum(widths)
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = str(h); shade(c, NAVY)
        for p in c.paragraphs:
            rtl(p, WD_ALIGN_PARAGRAPH.CENTER)
            for r in p.runs: font(r, size, True, "FFFFFF")
    for n, row in enumerate(rows):
        cells = t.add_row().cells
        vals = list(row[:len(headers)]) + [""] * max(0, len(headers)-len(row))
        for i, val in enumerate(vals):
            cells[i].text = str(val if val not in (None, "") else "—")
            if n % 2: shade(cells[i], PALE)
            for p in cells[i].paragraphs:
                rtl(p)
                p.paragraph_format.space_after = Pt(2)
                for r in p.runs: font(r, size)
    geometry(t, widths)
    keep_table_block(t)
    para(doc, "", 2)
    return t


def callout(doc, label, text, kind="info"):
    fills = {"info": LIGHT, "risk": "FCEBEC", "warn": "FFF6DD"}
    t = doc.add_table(rows=1, cols=1); t.style = "Table Grid"; shade(t.cell(0,0), fills[kind])
    p = t.cell(0,0).paragraphs[0]; rtl(p)
    font(p.add_run(label + ": "), 10, True, RED if kind == "risk" else NAVY)
    font(p.add_run(text), 10)
    geometry(t, [CONTENT_DXA])


def clean_inline(s):
    s = re.sub(r"!\[[^\]]*\]\([^)]*\)", "", s)
    s = re.sub(r"\[([^\]]+)\]\([^)]*\)", r"\1", s)
    return re.sub(r"[`*_]", "", s).strip()


def parse_table(lines, i):
    rows=[]
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        rows.append([clean_inline(x) for x in lines[i].strip().strip("|").split("|")]); i += 1
    if len(rows)>1 and all(re.fullmatch(r":?-{3,}:?", c.replace(" ","")) for c in rows[1]):
        return rows[0], rows[2:], i
    return None, None, i


SKIP_HEADINGS = re.compile(r"(الفجوات|المخاطر|غير متحقق|Coverage|الأدلة والملفات|نتيجة المهمة|حالة المستند|الحالة والنطاق)", re.I)


def add_markdown(doc, path, heading_shift=0, skip_local_audits=False, include_title=False):
    lines = path.read_text(encoding="utf-8-sig").splitlines(); i=0; in_code=False; code=[]; skip_level=None
    while i < len(lines):
        line=lines[i]
        hm=re.match(r"^(#{1,4})\s+(.*)$", line)
        if hm:
            lvl=len(hm.group(1)); title=clean_inline(hm.group(2))
            if skip_level and lvl <= skip_level: skip_level=None
            if skip_local_audits and SKIP_HEADINGS.search(title): skip_level=lvl; i+=1; continue
            if skip_level: i+=1; continue
        elif skip_level:
            i+=1; continue
        if line.startswith("```"):
            if in_code:
                p=doc.add_paragraph(style="Code RTL"); rtl(p, WD_ALIGN_PARAGRAPH.LEFT); p.paragraph_format.keep_together=False
                font(p.add_run("\n".join(code)), 7.0, name=MONO); code=[]; in_code=False
            else: in_code=True
            i+=1; continue
        if in_code: code.append(line); i+=1; continue
        if line.startswith("|"):
            h, rows, ni=parse_table(lines,i)
            if h is not None:
                table(doc,h,rows,size=7.2 if len(h)>4 else 7.8); i=ni; continue
        if hm:
            orig=len(hm.group(1))
            if orig==1 and not include_title: i+=1; continue
            heading(doc,clean_inline(hm.group(2)),min(3,max(1,orig+heading_shift)))
        elif re.match(r"^\s*[-*]\s+",line):
            p=doc.add_paragraph(style="List Bullet"); rtl(p); font(p.add_run(clean_inline(re.sub(r"^\s*[-*]\s+","",line))),9.8)
        elif re.match(r"^\s*\d+\.\s+",line):
            p=doc.add_paragraph(style="List Number"); rtl(p); font(p.add_run(clean_inline(re.sub(r"^\s*\d+\.\s+","",line))),9.8)
        elif line.strip().startswith(">"):
            callout(doc,"ملاحظة",clean_inline(line.strip().lstrip("> ")),"warn")
        elif line.strip(): para(doc,clean_inline(line),9.8)
        i+=1


def diagram_image(path, index):
    FIG_DIR.mkdir(parents=True, exist_ok=True)
    text=path.read_text(encoding="utf-8-sig")
    title=text.splitlines()[0].lstrip("# ")
    labels=[]
    for pat in (r'\["?([^\]"\n]{2,70})"?\]', r'\{"?([^}"\n]{2,70})"?\}', r'\("?([^\)"\n]{2,70})"?\)'):
        labels += re.findall(pat,text)
    labels=[clean_inline(x).replace("<br/>"," / ").replace("\\n"," / ") for x in labels]
    seen=[]
    for x in labels:
        if x not in seen and not re.fullmatch(r"[A-Za-z0-9_]+",x): seen.append(x)
    if not seen:
        seen=[clean_inline(x) for x in re.findall(r'^\s*[A-Za-z0-9_]+\s*:\s*(.+)$',text,re.M)]
    seen=seen[:12] or ["مخطط موثق في المصدر النصي"]
    img=Image.new("RGB",(1400,900),"white"); d=ImageDraw.Draw(img)
    try:
        ftitle=ImageFont.truetype("arialbd.ttf",38); fbox=ImageFont.truetype("arial.ttf",25); fsmall=ImageFont.truetype("arial.ttf",19)
    except OSError:
        ftitle=fbox=fsmall=ImageFont.load_default()
    d.rounded_rectangle((35,25,1365,875),radius=24,outline="#183B56",width=4,fill="#F8FBFC")
    d.text((700,65),title,font=ftitle,fill="#183B56",anchor="ma")
    cols=3 if len(seen)>6 else 2; rows=(len(seen)+cols-1)//cols
    boxw=390 if cols==3 else 570; boxh=min(120,max(78,650//max(rows,1)))
    xgap=(1260-cols*boxw)//max(cols-1,1); startx=70; starty=155
    for n,label in enumerate(seen):
        row,col=divmod(n,cols); x=startx+col*(boxw+xgap); y=starty+row*(boxh+34)
        d.rounded_rectangle((x,y,x+boxw,y+boxh),radius=16,fill="#EAF3F5",outline="#187B8C",width=3)
        wrapped="\n".join(textwrap.wrap(label,34 if cols==3 else 50)[:3])
        d.multiline_text((x+boxw/2,y+boxh/2),wrapped,font=fbox,fill="#183B56",anchor="mm",align="center",spacing=5)
        if n and col:
            d.line((x-25,y+boxh/2,x-5,y+boxh/2),fill="#9A6B00",width=5)
            d.polygon([(x-5,y+boxh/2),(x-17,y+boxh/2-8),(x-17,y+boxh/2+8)],fill="#9A6B00")
    d.text((700,845),"تمثيل طباعي مبسط مشتق من مخطط Mermaid الموثق",font=fsmall,fill="#58656D",anchor="mm")
    out=FIG_DIR/f"figure-{index:02d}.png"; img.save(out,dpi=(180,180)); return out,title


def diagram_content(path):
    text=path.read_text(encoding="utf-8-sig")
    title=text.splitlines()[0].lstrip("# ")
    labels=[]
    for pat in (r'\["?([^\]"\n]{2,70})"?\]', r'\{"?([^}"\n]{2,70})"?\}', r'\("?([^\)"\n]{2,70})"?\)'):
        labels += re.findall(pat,text)
    labels=[clean_inline(x).replace("<br/>"," / ") for x in labels]
    seen=[]
    for x in labels:
        if x not in seen and not re.fullmatch(r"[A-Za-z0-9_]+",x): seen.append(x)
    if not seen:
        seen=[clean_inline(x) for x in re.findall(r'^\s*[A-Za-z0-9_]+\s*:\s*(.+)$',text,re.M)]
    return title, (seen[:12] or ["مخطط موثق في المصدر النصي"])


def add_native_diagram(doc, path, index):
    title, labels = diagram_content(path)
    p=doc.add_paragraph(style="Caption"); rtl(p); font(p.add_run(f"شكل {index}: {title}"),9,True,GRAY)
    p=doc.add_paragraph(); rtl(p,WD_ALIGN_PARAGRAPH.CENTER); p.paragraph_format.keep_with_next=True; font(p.add_run(title),12,True,NAVY)
    rows=[]
    for i in range(0,len(labels),3):
        group=labels[i:i+3]
        while len(group)<3: group.append("")
        rows.append((group[0],"←",group[1],"←",group[2]))
    t=doc.add_table(rows=1,cols=5); t.style="Table Grid"
    widths=[2850,330,2850,330,2850]
    for i,val in enumerate(("المسار","","المرحلة","","النتيجة")):
        c=t.rows[0].cells[i]; c.text=val; shade(c,NAVY if i%2==0 else "FFFFFF")
        for pp in c.paragraphs:
            rtl(pp,WD_ALIGN_PARAGRAPH.CENTER)
            for rr in pp.runs: font(rr,8.2,True,"FFFFFF" if i%2==0 else GOLD)
    for n,row in enumerate(rows):
        cells=t.add_row().cells
        for i,val in enumerate(row):
            cells[i].text=val if val else "—"
            if i%2==0: shade(cells[i],LIGHT if n%2==0 else PALE)
            for pp in cells[i].paragraphs:
                rtl(pp,WD_ALIGN_PARAGRAPH.CENTER)
                pp.paragraph_format.space_after=Pt(1)
                for rr in pp.runs: font(rr,8.0,True if i%2 else False,GOLD if i%2 else NAVY)
    geometry(t,widths)
    keep_table_block(t, keep_whole=True)
    para(doc,"",2)
    return title


def add_figure(doc,path,caption):
    p=doc.add_paragraph(style="Caption"); rtl(p); font(p.add_run(caption),9,True,GRAY)
    p=doc.add_paragraph(); rtl(p,WD_ALIGN_PARAGRAPH.CENTER); r=p.add_run(); r.add_picture(str(path),width=Cm(16.0)); p.paragraph_format.keep_together=True
    doc_pr = r._r.xpath('.//wp:docPr')
    if doc_pr:
        doc_pr[0].set('descr', caption)


def extract_section(path, title_pattern):
    lines=path.read_text(encoding="utf-8-sig").splitlines(); out=[]; active=False; base=0
    for line in lines:
        m=re.match(r"^(#{1,4})\s+(.*)$",line)
        if m:
            lvl=len(m.group(1)); title=clean_inline(m.group(2))
            if active and lvl<=base: break
            if re.search(title_pattern,title,re.I): active=True; base=lvl; continue
        if active: out.append(line)
    return out


def main():
    progress=json.loads(PROGRESS.read_text(encoding="utf-8-sig")); now=datetime.now().astimezone()
    doc=Document(); setup(doc)
    # Formal editorial cover.
    para(doc,"SMARTFOUNDATION",11,True,TEAL,True)
    for _ in range(5): doc.add_paragraph()
    p=doc.add_paragraph(style="Title"); rtl(p,WD_ALIGN_PARAGRAPH.CENTER); font(p.add_run("توثيق نظام SmartFoundation"),28,True,NAVY)
    p=doc.add_paragraph(style="Subtitle"); rtl(p,WD_ALIGN_PARAGRAPH.CENTER); font(p.add_run("التوثيق الفني والوظيفي ودليل التشغيل والاستخدام"),15,True,GRAY)
    para(doc,"الإصدار 1.0.0",12,True,GOLD,True); para(doc,"5 سبتمبر 2026",11,False,GRAY,True)
    for _ in range(5): doc.add_paragraph()
    callout(doc,"حالة الوثيقة","وثيقة نهائية معتمدة؛ القيم الحساسة والأسرار وبيانات الاتصال غير مضمنة.","info")

    page_break(doc); heading(doc,"سجل الإصدارات",1)
    table(doc,["الإصدار","التاريخ","الحالة","نطاق التغيير"],[
        ("0.1","2026-08-22","مرحلي","توثيق المراحل حتى IncomeSystem والمصالحة الحية 7A"),
        ("1.0","2026-08-23","متحقق نهائياً","دمج المراحل 1–14، DATACORE الحية، Coverage Audit النهائي، ونسختي Word وPDF"),
        ("1.0.0","2026-09-05","محدث للإصدار","حالة التحقق الأمني النهائية، البرمجة الآمنة، المكونات والتراخيص، SBOM، الاستضافة، النسخ والاستعادة، ومعمارية المصادقة والجلسات")],
        [1100,1500,1600,CONTENT_DXA-4200],caption="جدول 1: سجل إصدارات وثيقة SmartFoundation")
    heading(doc,"نطاق التوثيق",1)
    table(doc,["التصنيف","البرامج"],[
        ("مشمولة","ControlPanel، Housing، IncomeSystem، ElectronicBillSystem، Home، Login"),
        ("مؤجلة","HousingCommandCenter، Maintenance، Support، Vehicle، Dashboard، Statistics، PopulationDensity، RealCharts"),
        ("أسطح مشتركة","Api، CrudController، ExportsController، NotificationsController، ReportsController، SessionController، SmartComponentController، Shared views")],
        [1900,CONTENT_DXA-1900],caption="جدول 2: البرامج والأسطح المشمولة والمؤجلة")
    callout(doc,"مصدر SQL النهائي","قاعدة DATACORE الحية هي المرجع النهائي لتعريفات SQL. SmartFoundation.Database Snapshot مرجع مقارنة فقط، ولا يُستخدم لإحياء ادعاء صححته المطابقة الحية.","warn")

    page_break(doc); heading(doc,"فهرس المحتويات",1)
    toc_entries=[
        ("1. الملخص التنفيذي","4"),("2. المعمارية","5"),("3. بنية الحل والنطاق","14"),
        ("4. الأمان والصلاحيات","19"),("5. ControlPanel","24"),("6. Housing","31"),
        ("7. IncomeSystem","49"),("8. ElectronicBillSystem","53"),("9. Home وLogin","58"),
        ("10. قاعدة البيانات الحية","62"),("11. Schema drift","71"),("12. Workflows والرسومات","72"),
        ("13. التقارير والتكاملات","80"),("14. بيئة التطوير والنشر","91"),("15. النسخ الاحتياطي والاستعادة","92"),
        ("16. استكشاف الأخطاء","93"),("17. دليل المطور","95"),("18. دليل المستخدم","100"),
        ("19. سجل الفجوات والمخاطر","106"),("20. Coverage Audit النهائي","115"),("21. حالة الإصدار 1.0.0 والأمن والتشغيل","119"),
        ("22. المصطلحات والاختصارات","124"),
    ]
    toc_rows=[]
    for i in range(0,len(toc_entries),2):
        left=toc_entries[i]; right=toc_entries[i+1] if i+1<len(toc_entries) else ("","")
        toc_rows.append((left[0],left[1],right[0],right[1]))
    table(doc,["الفصل","الصفحة","الفصل","الصفحة"],toc_rows,[3900,700,3900,700],8.2,caption="فهرس ثابت محدث وفق الرندر النهائي")

    chapters=[
        ("1. الملخص التنفيذي",None),
        ("2. المعمارية",DOC/"Work/03-System-Architecture-and-Shared-Components.md"),
        ("3. بنية الحل والنطاق",DOC/"Work/01-System-Inventory.md"),
        ("4. الأمان والصلاحيات",DOC/"Work/Security-Authentication-and-Authorization.md"),
        ("5. ControlPanel",DOC/"Work/04-ControlPanel-Program.md"),
        ("6. Housing",None),
        ("7. IncomeSystem",DOC/"Work/08-IncomeSystem.md"),
        ("8. ElectronicBillSystem",DOC/"Work/09-ElectronicBillSystem.md"),
        ("9. Home وLogin",DOC/"Work/10-Home-and-Login.md"),
        ("10. قاعدة البيانات الحية",DOC/"Work/11-Live-Database-and-ERD.md"),
        ("11. Schema drift",None),
        ("12. Workflows والرسومات",None),
        ("13. التقارير والتكاملات",None),
        ("14. بيئة التطوير والنشر",None),
        ("15. النسخ الاحتياطي والاستعادة",None),
        ("16. استكشاف الأخطاء",None),
        ("17. دليل المطور",DOC/"Developer-Guide.md"),
        ("18. دليل المستخدم",DOC/"User-Guide.md"),
        ("19. سجل الفجوات والمخاطر",DOC/"security-gap-status.md"),
        ("20. Coverage Audit النهائي",DOC/"Work/13-Final-Coverage-Audit.md"),
        ("21. حالة الإصدار 1.0.0 والأمن والتشغيل",DOC/"Work/14-Release-1.0.0-Readiness-Security-and-Operations.md"),
        ("22. المصطلحات والاختصارات",None),
    ]
    page_break(doc); heading(doc,chapters[0][0],1)
    para(doc,"SmartFoundation تطبيق ASP.NET Core MVC على .NET 8، ونقطة التشغيل الفعلية هي SmartFoundation.Mvc/Program.cs. يتكون المسار التشغيلي من MVC وUI وApplication وDataEngine، بينما يمثل مشروع SmartFoundation.Database لقطة مرجعية لا قاعدة التشغيل الحية.")
    para(doc,"النمط السائد، ولا سيما في Housing، يقرأ سياق Session ثم يمر عبر MastersServies وDataSet وبوابة Masters_DataLoad، ويكتب عبر CrudController وMasters_CRUD. تبني Controllers تكوين SmartRenderer على الخادم وتبقي Razor Views رقيقة.")
    callout(doc,"نتيجة المصالحة","اكتملت المحادثات 1–13 وLive Database Reconciliation. تم التحقق من 274 كائن SQL حي ضمن النطاق: 118 جدولًا + 44 View + 76 Stored Procedure + 34 Function + 2 Trigger.","info")
    callout(doc,"حد الاعتماد","التغطية تعني اكتمال الأسطح المكتشفة في المستودع وقراءة catalog الحية؛ لا تعني تنفيذ إجراءات الأعمال أو اختبار حسابات فعلية أو تحقق E2E.","warn")

    for title,path in chapters[1:10]:
        page_break(doc); heading(doc,title,1)
        if path: add_markdown(doc,path,heading_shift=0,skip_local_audits=True)
        if title.startswith("6."):
            for f,sub in (("05-Housing-Definitions.md","تعريفات Housing"),("06-Housing-Procedures.md","إجراءات Housing"),("07-Housing-Waiting-Lists-and-Imports.md","الانتظار والإسناد والاستيراد")):
                heading(doc,sub,2); add_markdown(doc,DOC/"Work"/f,heading_shift=1,skip_local_audits=True)

    page_break(doc); heading(doc,"11. Schema drift",1)
    drift=extract_section(DOC/"Work/11-Live-Database-and-ERD.md",r"Schema drift")
    tmp=DOC/".word_qa"/"schema_drift.md"; tmp.parent.mkdir(parents=True,exist_ok=True); tmp.write_text("\n".join(drift),encoding="utf-8")
    add_markdown(doc,tmp,heading_shift=-1,include_title=True)
    callout(doc,"قاعدة القراءة","يعرض هذا الفصل الاختلاف التاريخي فقط. عند التعارض، وصف Live هو السلوك الحالي، وSnapshot لا يثبت الحالة التشغيلية.","warn")

    page_break(doc); heading(doc,"12. Workflows والرسومات",1)
    para(doc,"تمت مراجعة مخططات Mermaid المصدرية وتحويلها إلى تمثيلات طباعية مبسطة بعرض ثابت ملائم لـA4. تحتفظ ملفات المصدر بالتفاصيل الكاملة، بينما تركز الأشكال هنا على العقد الأساسية وتدفقها.")
    diagrams=sorted((DOC/"Diagrams").glob("*.md"))
    for idx,path in enumerate(diagrams,1):
        add_native_diagram(doc,path,idx)

    ops=DOC/"Work/12-Reports-Integrations-and-Operations.md"
    for chnum,title,pattern in ((13,"التقارير والتكاملات",r"بنية التقارير|التكاملات"),(14,"بيئة التطوير والنشر",r"بيئة التطوير|النشر على IIS"),(15,"النسخ الاحتياطي والاستعادة",r"النسخ الاحتياطي"),(16,"استكشاف الأخطاء",r"استكشاف الأخطاء")):
        page_break(doc); heading(doc,f"{chnum}. {title}",1)
        # Include the full operations document only once, then concise extracted sections.
        if chnum==13: add_markdown(doc,ops,heading_shift=0,skip_local_audits=True)
        else:
            lines=extract_section(ops,pattern); t=DOC/".word_qa"/f"ops_{chnum}.md"; t.write_text("\n".join(lines),encoding="utf-8"); add_markdown(doc,t,heading_shift=-1,include_title=True)

    for title,path in chapters[16:18]:
        page_break(doc); heading(doc,title,1); add_markdown(doc,path,heading_shift=0,skip_local_audits=True)

    page_break(doc); heading(doc,"19. سجل الفجوات والمخاطر",1)
    add_markdown(doc,DOC/"security-gap-status.md",heading_shift=0,skip_local_audits=False)

    page_break(doc); heading(doc,"20. Coverage Audit النهائي",1)
    add_markdown(doc,DOC/"Work/13-Final-Coverage-Audit.md",heading_shift=0,skip_local_audits=False)

    page_break(doc); heading(doc,chapters[20][0],1)
    add_markdown(doc,chapters[20][1],heading_shift=0,skip_local_audits=False)

    page_break(doc); heading(doc,"22. المصطلحات والاختصارات",1)
    terms=[
        ("DATACORE","قاعدة SQL Server الحية والمصدر النهائي لتعريفات SQL ضمن هذا التوثيق."),
        ("Snapshot","مشروع SmartFoundation.Database المرجعي؛ يستخدم للمقارنة وفهم النية ولا يثبت الحالة الحية."),
        ("MVC","طبقة الويب التي تضم Program وControllers وViews ونقطة التشغيل الفعلية."),
        ("UI / SmartRenderer","مكونات العرض المشتركة التي تفسر SmartPageViewModel وتبني الجداول والنماذج."),
        ("MastersServies","خدمة بوابة تطبيقية نشطة للقراءة والـCRUD في نمط Housing."),
        ("Masters_DataLoad","Stored Procedure بوابة للقراءة، توجه بحسب pageName_."),
        ("Masters_CRUD","Stored Procedure بوابة للكتابة، توجه بحسب pageName_ وActionType وتطبق فحوصاً مشتركة."),
        ("DataSet","حاوية result sets؛ الجدول الأول غالباً للصلاحيات والجداول اللاحقة لبيانات الصفحة والـDDL."),
        ("DDL","قوائم الاختيار والبيانات المرجعية التي تغذي الحقول."),
        ("Idara","السياق الإداري الذي يقيد بيانات وعمليات كثيرة."),
        ("E2E","اختبار End-to-End من واجهة المستخدم إلى قاعدة البيانات والعودة."),
        ("Schema drift","اختلاف بنيوي أو وظيفي بين DATACORE الحية والـSnapshot."),
        ("RPO/RTO","أهداف نقطة الاستعادة ووقت الاستعادة في استمرارية الأعمال."),
        ("SBOM","قائمة جرد للمكونات والتبعيات بصيغة معيارية؛ ليست اختبار ثغرات مستقلاً."),
        ("SAST / SCA / DAST","فحص الشفرة الثابت، وتحليل المكونات، والفحص الديناميكي ضمن نطاقات التحقق المحددة."),
        ("IAST","اختبار أمان تفاعلي لم ينفذ ضمن نطاق فريق التطوير للإصدار 1.0.0."),
    ]
    table(doc,["المصطلح","التعريف الموحد"],terms,[2200,CONTENT_DXA-2200],8.2,caption="جدول 5: المصطلحات والاختصارات")

    doc.core_properties.title="توثيق نظام SmartFoundation - التوثيق الفني والوظيفي ودليل التشغيل والاستخدام"
    doc.core_properties.subject="التوثيق الفني والوظيفي والتشغيلي لنظام SmartFoundation"
    doc.core_properties.author="sami alamri"
    doc.core_properties.version="1.0.0"
    doc.core_properties.keywords="SmartFoundation, DATACORE, Housing, ControlPanel, IncomeSystem, ElectronicBillSystem"
    doc.core_properties.comments="لا تحتوي الوثيقة على أسرار اتصال أو قيم كلمات مرور أو بيانات مستخدمين فعلية."
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
