from docx import Document

doc = Document()
doc.add_heading("Render Probe", 0)
doc.add_paragraph("SmartFoundation visual verification probe.")
doc.save("Documentation/.interim_qa/render-probe.docx")
