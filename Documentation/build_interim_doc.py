from __future__ import annotations

import json
import re
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
DOC_DIR = ROOT / "Documentation"
OUT = DOC_DIR / "SmartFoundation_Interim_Documentation.docx"
PROGRESS = DOC_DIR / "Documentation-Progress.json"

FONT = "Arial"
NAVY = "183B56"
TEAL = "187B8C"
LIGHT = "EAF3F5"
PALE = "F4F7F9"
GRAY = "5F6B73"
RED = "A33A3A"
GOLD = "9A6B00"
PAGE_WIDTH_DXA = 11906
MARGIN_DXA = 1020
CONTENT_DXA = PAGE_WIDTH_DXA - 2 * MARGIN_DXA


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=120, bottom=90, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for tag, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{tag}"))
        if node is None:
            node = OxmlElement(f"w:{tag}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def rtl_paragraph(p, align=WD_ALIGN_PARAGRAPH.RIGHT):
    p.alignment = align
    p_pr = p._p.get_or_add_pPr()
    bidi = p_pr.find(qn("w:bidi"))
    if bidi is None:
        bidi = OxmlElement("w:bidi")
        p_pr.append(bidi)
    bidi.set(qn("w:val"), "1")
    return p


def set_run(run, size=None, bold=None, color=None, italic=None, font=FONT):
    run.font.name = font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:cs"), font)
    run._element.get_or_add_rPr().append(OxmlElement("w:rtl"))
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_field(run, instruction, display=""):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instruction
    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    text = OxmlElement("w:t")
    text.text = display
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for node in (begin, instr, sep, text, end):
        run._r.append(node)


def configure_styles(doc):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = FONT
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:cs"), FONT)
    normal.paragraph_format.space_after = Pt(5)
    normal.paragraph_format.line_spacing = 1.18
    for name, size, color, before, after in (
        ("Title", 27, NAVY, 0, 10),
        ("Subtitle", 14, GRAY, 0, 10),
        ("Heading 1", 18, NAVY, 14, 7),
        ("Heading 2", 14, TEAL, 11, 5),
        ("Heading 3", 12, NAVY, 8, 4),
    ):
        s = styles[name]
        s.font.name = FONT
        s.font.size = Pt(size)
        s.font.color.rgb = RGBColor.from_string(color)
        s.font.bold = name != "Subtitle"
        s._element.rPr.rFonts.set(qn("w:cs"), FONT)
        s.paragraph_format.space_before = Pt(before)
        s.paragraph_format.space_after = Pt(after)
        s.paragraph_format.keep_with_next = True
    if "Arabic Code" not in styles:
        code = styles.add_style("Arabic Code", WD_STYLE_TYPE.PARAGRAPH)
        code.font.name = "Courier New"
        code.font.size = Pt(7.5)
        code.paragraph_format.space_after = Pt(4)
        code.paragraph_format.line_spacing = 1.0
    if "Figure Caption" not in styles:
        cap = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
        cap.font.name = FONT
        cap.font.size = Pt(9)
        cap.font.bold = True
        cap.font.color.rgb = RGBColor.from_string(GRAY)
        cap.paragraph_format.space_before = Pt(4)
        cap.paragraph_format.space_after = Pt(6)
        cap.paragraph_format.keep_with_next = True


def configure_section(section, title="نسخة مرحلية — غير نهائية"):
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.7)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)
    section.header_distance = Cm(0.7)
    section.footer_distance = Cm(0.7)
    hp = section.header.paragraphs[0]
    rtl_paragraph(hp)
    hp.text = ""
    r = hp.add_run(title)
    set_run(r, 8.5, True, RED)
    fp = section.footer.paragraphs[0]
    rtl_paragraph(fp, WD_ALIGN_PARAGRAPH.CENTER)
    fp.text = ""
    r = fp.add_run("SmartFoundation  |  نسخة مرحلية — غير نهائية  |  صفحة ")
    set_run(r, 8, False, GRAY)
    add_field(r, " PAGE ", "1")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl_pr = table._tbl.tblPr
    layout = tbl_pr.find(qn("w:tblLayout"))
    if layout is None:
        layout = OxmlElement("w:tblLayout")
        tbl_pr.append(layout)
    layout.set(qn("w:type"), "fixed")
    tbl_w = tbl_pr.find(qn("w:tblW"))
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.find(qn("w:tcW"))
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:w"), str(widths[i]))
            tc_w.set(qn("w:type"), "dxa")
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            for p in cell.paragraphs:
                rtl_paragraph(p)
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    set_run(run, 8.5)
    first = table.rows[0]._tr.get_or_add_trPr()
    rep = OxmlElement("w:tblHeader")
    rep.set(qn("w:val"), "true")
    first.append(rep)


def add_table(doc, headers, rows, widths=None, font_size=8.2):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    if widths is None:
        widths = [CONTENT_DXA // len(headers)] * len(headers)
        widths[-1] += CONTENT_DXA - sum(widths)
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = str(h)
        set_cell_shading(cell, NAVY)
        for p in cell.paragraphs:
            rtl_paragraph(p, WD_ALIGN_PARAGRAPH.CENTER)
            for run in p.runs:
                set_run(run, font_size, True, "FFFFFF")
    for n, row in enumerate(rows):
        cells = table.add_row().cells
        normalized = list(row[:len(headers)]) + [""] * max(0, len(headers) - len(row))
        for i, value in enumerate(normalized):
            cells[i].text = str(value if value is not None else "—")
            if n % 2:
                set_cell_shading(cells[i], PALE)
            for p in cells[i].paragraphs:
                rtl_paragraph(p)
                for run in p.runs:
                    set_run(run, font_size)
    set_table_geometry(table, widths)
    doc.add_paragraph().paragraph_format.space_after = Pt(1)
    return table


def add_heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    rtl_paragraph(p)
    r = p.add_run(text)
    set_run(r, {1:18, 2:14, 3:12}[level], True, {1:NAVY, 2:TEAL, 3:NAVY}[level])
    return p


def add_para(doc, text="", bold=False, color=None, size=10.5, style=None, center=False):
    p = doc.add_paragraph(style=style)
    rtl_paragraph(p, WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.RIGHT)
    r = p.add_run(text)
    set_run(r, size, bold, color)
    return p


def add_callout(doc, label, text, severity="info"):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    fill = {"info": LIGHT, "risk": "FCEBEC", "warn": "FFF6DD"}.get(severity, LIGHT)
    set_cell_shading(table.cell(0, 0), fill)
    p = table.cell(0, 0).paragraphs[0]
    rtl_paragraph(p)
    r = p.add_run(f"{label}: ")
    set_run(r, 10, True, RED if severity == "risk" else NAVY)
    r = p.add_run(text)
    set_run(r, 10)
    set_table_geometry(table, [CONTENT_DXA])


def strip_inline(text):
    return re.sub(r"[`*_]", "", text).strip()


def markdown_table(lines, start):
    rows = []
    i = start
    while i < len(lines) and lines[i].lstrip().startswith("|"):
        rows.append([strip_inline(x) for x in lines[i].strip().strip("|").split("|")])
        i += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", c.replace(" ", "")) for c in rows[1]):
        return rows[0], rows[2:], i
    return None, None, start


def add_markdown(doc, path, heading_shift=0, include_title=True):
    text = path.read_text(encoding="utf-8-sig")
    lines = text.splitlines()
    i = 0
    in_code = False
    code = []
    while i < len(lines):
        line = lines[i]
        if line.startswith("```"):
            if in_code:
                p = doc.add_paragraph(style="Arabic Code")
                rtl_paragraph(p, WD_ALIGN_PARAGRAPH.LEFT)
                p.paragraph_format.keep_together = False
                r = p.add_run("\n".join(code))
                set_run(r, 7.2, font="Courier New")
                code = []
                in_code = False
            else:
                in_code = True
            i += 1
            continue
        if in_code:
            code.append(line)
            i += 1
            continue
        if line.startswith("|") and i + 1 < len(lines):
            h, rows, new_i = markdown_table(lines, i)
            if h is not None:
                widths = [CONTENT_DXA // len(h)] * len(h)
                widths[-1] += CONTENT_DXA - sum(widths)
                add_table(doc, h, rows, widths, 7.6 if len(h) > 4 else 8.2)
                i = new_i
                continue
        m = re.match(r"^(#{1,4})\s+(.*)$", line)
        if m:
            original = len(m.group(1))
            if original == 1 and not include_title:
                i += 1
                continue
            level = min(3, max(1, original + heading_shift))
            add_heading(doc, strip_inline(m.group(2)), level)
        elif re.match(r"^\s*[-*]\s+", line):
            text_value = strip_inline(re.sub(r"^\s*[-*]\s+", "", line))
            p = doc.add_paragraph(style="List Bullet")
            rtl_paragraph(p)
            r = p.add_run(text_value)
            set_run(r, 10)
        elif re.match(r"^\s*\d+\.\s+", line):
            text_value = strip_inline(re.sub(r"^\s*\d+\.\s+", "", line))
            p = doc.add_paragraph(style="List Number")
            rtl_paragraph(p)
            r = p.add_run(text_value)
            set_run(r, 10)
        elif line.strip():
            add_para(doc, strip_inline(line), size=10)
        i += 1


def chapter_break(doc):
    doc.add_page_break()


def main():
    progress = json.loads(PROGRESS.read_text(encoding="utf-8-sig"))
    generated = datetime.now().astimezone()
    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])

    # Cover: editorial_cover pattern, adapted to Arabic/RTL and A4.
    add_para(doc, "SMARTFOUNDATION", bold=True, color=TEAL, size=11, center=True)
    for _ in range(4):
        doc.add_paragraph()
    p = doc.add_paragraph(style="Title")
    rtl_paragraph(p, WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("توثيق نظام SmartFoundation")
    set_run(r, 27, True, NAVY)
    p = doc.add_paragraph(style="Subtitle")
    rtl_paragraph(p, WD_ALIGN_PARAGRAPH.CENTER)
    r = p.add_run("نسخة مرحلية للمراجعة")
    set_run(r, 17, True, RED)
    add_para(doc, "نسخة مرحلية — غير نهائية", bold=True, color=RED, size=12, center=True)
    doc.add_paragraph()
    add_para(doc, f"الإصدار المرحلي: {progress.get('documentationVersion', '1.0')}-interim", bold=True, center=True)
    add_para(doc, f"تاريخ الإنشاء: {generated.strftime('%Y-%m-%d %H:%M %Z')}", center=True)
    add_para(doc, "مبنية حصراً على المراحل المسجلة completed حتى المحادثة 8، مع مصالحة قاعدة البيانات الحية 7A.", center=True)
    for _ in range(4):
        doc.add_paragraph()
    add_callout(doc, "تنبيه", "هذه نسخة استعراضية مرحلية، ولا تستبدل الوثيقة النهائية ولا تعني اكتمال مشروع التوثيق.", "risk")

    chapter_break(doc)
    add_heading(doc, "فهرس المحتويات", 1)
    toc_items = [
        "تعريف النسخة المرحلية ونطاقها",
        "حالة المراحل والمحادثات",
        "الملخص التنفيذي الحالي",
        "Coverage المرحلي",
        "طبقات ومشاريع النظام",
        "جرد النظام ونطاق التوثيق",
        "المعمارية والمكونات المشتركة",
        "الأمان والمصادقة والترخيص والجلسات",
        "ControlPanel",
        "Housing Definitions",
        "Housing Procedures",
        "Housing Waiting Lists and Imports",
        "Live Database Reconciliation",
        "IncomeSystem",
        "الرسومات المنجزة",
        "أدلة المستخدم المكتملة",
        "سجل الفجوات والمخاطر الحالية",
        "المراحل والمحادثات المتبقية",
    ]
    for item in toc_items:
        p = doc.add_paragraph(style="List Bullet")
        rtl_paragraph(p)
        r = p.add_run(item)
        set_run(r, 10.5)

    chapter_break(doc)
    add_heading(doc, "1. تعريف النسخة المرحلية ونطاقها", 1)
    add_callout(doc, "حالة الوثيقة", "قيد العمل. لا تتضمن المراحل غير المكتملة ولا تنشئ محتوى بديلاً للفجوات.", "warn")
    add_heading(doc, "مصادر الحقيقة", 2)
    for item in (
        "الكود النشط لسلوك التطبيق.",
        "قاعدة DATACORE الحية لتعريفات SQL.",
        "نتائج Live Database Reconciliation.",
        "SmartFoundation.Database للمقارنة المرجعية فقط.",
    ):
        p = doc.add_paragraph(style="List Number")
        rtl_paragraph(p)
        r = p.add_run(item)
        set_run(r, 10)
    add_callout(doc, "قاعدة التنبيه", "أي معلومة ما زالت مبنية على Snapshot دون تحقق حي تحمل عبارة: غير متحقق منها من قاعدة البيانات الحية.", "warn")
    add_heading(doc, "النطاق الحالي", 2)
    add_para(doc, "تشمل النسخة نتائج الجرد والمعمارية والأمان وControlPanel وأجزاء Housing المكتملة وIncomeSystem ومصالحة DATACORE الحية حتى الآن، إضافة إلى الرسومات وأدلة المستخدم المكتملة.")
    add_para(doc, "لا تشمل توثيق ElectronicBillSystem أو Home/Login التفصيلي أو قاعدة البيانات الشاملة أو النشر والتشغيل أو النسخة النهائية أو PDF النهائي.")

    add_heading(doc, "2. حالة المراحل والمحادثات", 1)
    rows = []
    outputs = {
        "inventory_and_scope": "Work/01-System-Inventory.md",
        "architecture_and_shared_components": "Work/03-System-Architecture-and-Shared-Components.md + Diagrams/01–03",
        "security_authentication_and_authorization": "Work/Security-* + رسمان أمنيان",
        "control_panel": "Work/04-* + UserManual/ControlPanel + Diagram/04",
        "housing_definitions": "Work/05-* + Diagrams/05–06",
        "housing_procedures": "Work/06-* + UserManual + Diagrams/07–09",
        "housing_waiting_lists_and_imports": "Work/07-* + UserManual + Diagrams/10–12",
        "live_database_reconciliation_for_conversations_1_to_7": "Work/07A-* + تحديث الوثائق السابقة",
        "income_system": "Work/08-* + UserManual/IncomeSystem + Diagrams/13–14",
    }
    for st in progress["completedStages"]:
        rows.append((st["conversation"], st["stage"], st["status"], outputs.get(st["stage"], "موثق في Documentation/")))
    n = progress["nextStage"]
    rows.append((n["conversation"], n["stage"], n["status"], "لم يُنشأ محتوى مرحلي"))
    add_table(doc, ["المحادثة", "المرحلة", "الحالة", "الملفات الناتجة"], rows, [850, 2650, 1200, CONTENT_DXA-4700], 7.8)
    add_heading(doc, "التصنيف المرحلي", 2)
    add_table(doc, ["التصنيف", "المراحل"], [
        ("المكتمل", "المحادثات 1، 2، 3، 4، 5، 6، 7، 7A، 8"),
        ("الجاري", "مشروع التوثيق العام قيد العمل؛ لا توجد مرحلة برنامج مسجلة in_progress في السجل"),
        ("المتبقي", "المحادثة 9 وما يليها حسب خطة التنفيذ"),
    ], [1700, CONTENT_DXA-1700])

    add_heading(doc, "3. الملخص التنفيذي الحالي", 1)
    add_para(doc, "SmartFoundation تطبيق ASP.NET Core MVC على .NET 8. نقطة التشغيل الفعلية هي SmartFoundation.Mvc/Program.cs، ويتبع المسار السائد أربع طبقات تشغيلية: MVC وUI وApplication وDataEngine. يعتمد نمط Housing المرجعي على MastersServies وDataSet وبوابتي Masters_DataLoad وMasters_CRUD، مع واجهات Razor رقيقة يبني Controller تكوينها على الخادم.")
    add_para(doc, "أثبتت المراحل المكتملة توثيق ControlPanel وثلاث مجموعات Housing وIncomeSystem، مع مطابقة حية اختيارية القراءة لتعريفات SQL في DATACORE. بقيت فجوات تشغيلية وأمنية مهمة، ولا تعني تغطية أسطح المستودع اكتمال اختبارات E2E أو تنفيذ الإجراءات التجارية.")
    add_callout(doc, "أهم مخاطرة", "سياق CRUD الأمني (entrydata وidaraID وpageName_ وActionType) يصل من العميل في مسارات مشتركة، بينما تستخدمه SQL في فحص الصلاحية والتدقيق؛ لم يُنفذ اختبار استغلال أو كتابة.", "risk")

    add_heading(doc, "4. Coverage المرحلي", 1)
    cov_rows = [
        ("Programs", "3 موثقة تفصيلياً من 6 مشمولة", "50.0%", "ControlPanel وHousing وIncomeSystem؛ المستبعدة لا تدخل النسبة"),
        ("Controller files", "29 من 38", "76.3%", "القيمة المباشرة من coverage"),
        ("Logical Controllers", "3 من 8 حسب الحقل الإجمالي", "37.5%", "مع ملاحظة partial controllers"),
        ("Actions", "35 عملية MVC موثقة في الوحدات المكتملة", "لا يتوفر مقام شامل", "4 + 6 + 4 + 13 + 8"),
        ("Views", "28 من 35", "80.0%", "القيمة المباشرة من coverage"),
        ("Services", "المكونات المشتركة موثقة مع خدمات الوحدات المكتملة", "لا يتوفر عدّ شامل", "لا يُخترع مقام غير مسجل"),
        ("Stored Procedures", "55 من 62 محددة أولياً", "88.7%", "تعريف/استخدام موثق، لا يعني تنفيذها"),
        ("SQL Views", "19 متحققة في 7A + 9 مرتبطة بـ IncomeSystem", "لا يتوفر مقام شامل", "قد يوجد تداخل؛ لا تجمع كنسبة نهائية"),
        ("Tables", "موثقة ضمن الفصول حسب علاقات الصفحات", "لا يتوفر عدّ شامل", "تحليل قاعدة البيانات الشامل pending"),
        ("Functions", "19 Views/Functions متحققة في 7A + 9 مرتبطة بـ IncomeSystem", "لا يتوفر مقام منفصل", "السجل لا يفصل العددين عالمياً"),
        ("Permissions", "72 تسمية موثقة في Housing وIncomeSystem", "لا يتوفر مقام شامل", "15 + 20 + 26 + 11؛ لا يشمل تقديراً لـ ControlPanel"),
        ("Workflows", "مسارات مشتركة و9 مراحل مكتملة ممثلة", "لا يتوفر مقام شامل", "لا توجد قائمة workflow عالمية"),
        ("User Manual pages", "4 ملفات دليل مكتملة", "4/4 للبرامج الموثقة ذات الدليل", "ControlPanel، Housing Procedures، Waiting/Imports، IncomeSystem"),
    ]
    add_table(doc, ["البند", "التغطية الحالية", "النسبة", "ملاحظة"], cov_rows, [1500, 2550, 1250, CONTENT_DXA-5300], 7.4)

    add_heading(doc, "5. طبقات ومشاريع النظام", 1)
    add_table(doc, ["المشروع", "الحالة", "الدور"], [
        ("SmartFoundation.Mvc", "موثق معمارياً", "التشغيل والعرض وComposition Root وControllers وViews"),
        ("SmartFoundation.UI", "موثق معمارياً", "ViewComponents وViewModels ومكونات العرض المشتركة"),
        ("SmartFoundation.Application", "موثق معمارياً", "خدمات التطبيق وProcedureMapper وبوابة MastersServies"),
        ("SmartFoundation.DataEngine", "موثق معمارياً", "Dapper وSmartComponentService وConnectionFactory"),
        ("SmartFoundation.Database", "مرجع SQL فقط", "Snapshot للمقارنة؛ DATACORE الحية هي المصدر النهائي لتعريفات SQL"),
    ], [2400, 1700, CONTENT_DXA-4100])
    add_callout(doc, "قاعدة معمارية", "الطبقات التشغيلية أربع. SmartFoundation.Database ليس طبقة تشغيلية ولا يثبت حالة القاعدة الحية.", "info")

    # Full completed work products.
    work_files = [
        "01-System-Inventory.md",
        "03-System-Architecture-and-Shared-Components.md",
        "Security-Authentication-and-Authorization.md",
        "04-ControlPanel-Program.md",
        "05-Housing-Definitions.md",
        "06-Housing-Procedures.md",
        "07-Housing-Waiting-Lists-and-Imports.md",
        "07A-Live-Database-Reconciliation.md",
        "08-IncomeSystem.md",
    ]
    chapter_no = 6
    for name in work_files:
        chapter_break(doc)
        title = (DOC_DIR / "Work" / name).read_text(encoding="utf-8-sig").splitlines()[0].lstrip("# ")
        add_heading(doc, f"{chapter_no}. {title}", 1)
        add_markdown(doc, DOC_DIR / "Work" / name, heading_shift=1, include_title=False)
        chapter_no += 1

    chapter_break(doc)
    add_heading(doc, f"{chapter_no}. الرسومات المنجزة", 1)
    add_para(doc, "تُدرج الرسومات المكتملة كما هي كمخططات Mermaid موثقة مع Caption. لم تُنشأ رسومات جديدة لسد أي فجوة.")
    diagram_files = sorted((DOC_DIR / "Diagrams").glob("*.md"))
    for idx, path in enumerate(diagram_files, 1):
        title = path.read_text(encoding="utf-8-sig").splitlines()[0].lstrip("# ")
        p = doc.add_paragraph(style="Figure Caption")
        rtl_paragraph(p)
        r = p.add_run(f"شكل {idx}: {title}")
        set_run(r, 9, True, GRAY)
        add_markdown(doc, path, heading_shift=2, include_title=False)
    chapter_no += 1

    chapter_break(doc)
    add_heading(doc, f"{chapter_no}. أدلة المستخدم المكتملة", 1)
    manuals = sorted((DOC_DIR / "UserManual").glob("*.md"))
    for path in manuals:
        title = path.read_text(encoding="utf-8-sig").splitlines()[0].lstrip("# ")
        add_heading(doc, title, 2)
        add_markdown(doc, path, heading_shift=2, include_title=False)
    chapter_no += 1

    chapter_break(doc)
    add_heading(doc, f"{chapter_no}. سجل الفجوات والمخاطر الحالية", 1)
    add_markdown(doc, DOC_DIR / "security-gap-status.md", heading_shift=1, include_title=False)

    add_heading(doc, f"{chapter_no + 1}. المراحل والمحادثات المتبقية", 1)
    remaining = [
        ("9", "ElectronicBillSystem", "pending"),
        ("10", "Home وLogin", "متبقية وفق خطة التنفيذ"),
        ("11", "تحليل قاعدة البيانات وتتبع الجداول والـ Views والـ Functions والـ Triggers", "pending"),
        ("12", "استكمال ERD والرسومات وWorkflows اللازمة بعد اكتمال التحليل", "pending"),
        ("13", "التقارير والتكاملات والتشغيل والنشر والنسخ الاحتياطي واستكشاف الأخطاء", "pending"),
        ("14", "دليل المطور واستكمال دليل المستخدم", "pending"),
        ("15", "تجميع Word النهائي وتدقيقه", "pending"),
        ("16", "PDF النهائي وCoverage Audit والتحقق النهائي", "pending"),
    ]
    add_table(doc, ["المحادثة/الترتيب", "المرحلة", "الحالة"], remaining, [1450, CONTENT_DXA-2850, 1400])
    add_callout(doc, "خاتمة مرحلية", "لا تغيّر هذه النسخة lastCompletedStage أو nextStage أو الحالة العامة للمشروع، ولا تستبدل SmartFoundation_System_Documentation.docx.", "risk")

    doc.core_properties.title = "توثيق SmartFoundation — نسخة مرحلية للمراجعة"
    doc.core_properties.subject = "Interim documentation based on completed stages only"
    doc.core_properties.author = "SmartFoundation Documentation Project"
    doc.core_properties.comments = "نسخة مرحلية — غير نهائية"
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    main()
